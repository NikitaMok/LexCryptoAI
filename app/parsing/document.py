from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass
from enum import Enum
from pathlib import Path

DOCX_SUFFIXES = {".docx"}
PDF_SUFFIXES = {".pdf"}

# Мягкий перенос и неразрывный пробел встречаются в контрактах из Word и мешают
# и регулярным выражениям, и эмбеддингам.
_INVISIBLE = str.maketrans({"\u00ad": "", "\u200b": "", "\u00a0": " ", "\ufeff": ""})

_HYPHEN_BREAK = re.compile(r"(\w)[-\u2010\u2011]\n(\w)")
_MULTISPACE = re.compile(r"[ \t\u2000-\u200a]+")
_BLANK_RUN = re.compile(r"\n{3,}")


class SourceFormat(str, Enum):
    DOCX = "docx"
    PDF = "pdf"


class UnsupportedFormatError(ValueError):
    pass


class EmptyDocumentError(ValueError):
    pass


@dataclass(frozen=True)
class Document:
    source_format: SourceFormat
    paragraphs: list[str]
    page_count: int | None = None

    @property
    def text(self) -> str:
        return "\n".join(self.paragraphs)


def normalize(raw: str) -> str:
    text = unicodedata.normalize("NFKC", raw).translate(_INVISIBLE)
    text = _HYPHEN_BREAK.sub(r"\1\2", text)
    text = _MULTISPACE.sub(" ", text)
    return _BLANK_RUN.sub("\n\n", text).strip()


def _split_paragraphs(text: str) -> list[str]:
    return [line.strip() for line in text.split("\n") if line.strip()]


def _load_docx(path: Path) -> Document:
    from docx import Document as DocxDocument

    document = DocxDocument(str(path))
    blocks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            blocks.extend(cell.text for cell in row.cells)

    return Document(
        source_format=SourceFormat.DOCX,
        paragraphs=_split_paragraphs(normalize("\n".join(blocks))),
    )


def _load_pdf(path: Path) -> Document:
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")

    return Document(
        source_format=SourceFormat.PDF,
        paragraphs=_split_paragraphs(normalize("\n".join(pages))),
        page_count=len(pages),
    )


def load_document(path: str | Path) -> Document:
    source = Path(path)
    if not source.is_file():
        raise FileNotFoundError(source)

    suffix = source.suffix.lower()
    if suffix in DOCX_SUFFIXES:
        document = _load_docx(source)
    elif suffix in PDF_SUFFIXES:
        document = _load_pdf(source)
    else:
        raise UnsupportedFormatError(
            f"поддерживаются только PDF и DOCX, получен «{suffix or source.name}»"
        )

    if not document.paragraphs:
        raise EmptyDocumentError(
            f"из документа {source.name} не извлечён текст: "
            "вероятно, это скан без распознанного слоя"
        )

    return document
