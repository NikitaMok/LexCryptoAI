from datetime import date
from pathlib import Path

import pytest

from app.rules.contract import ContractView
from app.rules.engine import ContractStatus, FindingStatus, evaluate
from app.rules.predicates import registered_names
from app.rules.registry import Severity, load_rules

LAW_IN_FORCE = date(2026, 9, 1)


@pytest.fixture(scope="module")
def rules():
    return load_rules()


@pytest.fixture(scope="module")
def compliant(compliant_docx: Path):
    return ContractView.from_file(compliant_docx)


@pytest.fixture(scope="module")
def violating(violating_docx: Path):
    return ContractView.from_file(violating_docx)


class TestWiring:
    def test_every_declared_check_is_implemented(self, rules):
        declared = {rule.check.predicate for rule in rules.automated()}

        assert declared <= registered_names(), (
            f"в конфиге объявлены несуществующие проверки: {sorted(declared - registered_names())}"
        )

    def test_no_orphan_predicates(self, rules):
        declared = {rule.check.predicate for rule in rules.automated()}

        assert registered_names() <= declared, (
            f"проверки реализованы, но не подключены к правилам: "
            f"{sorted(registered_names() - declared)}"
        )

    def test_most_mandatory_rules_are_automated(self, rules):
        mandatory = rules.mandatory()
        automated = [rule for rule in mandatory if rule.check is not None]

        assert len(automated) >= len(mandatory) - 3, (
            "обязательных правил без автоматической проверки больше трёх: "
            f"{[rule.code for rule in mandatory if rule.check is None]}"
        )


class TestCompliantContract:
    def test_status_is_green(self, compliant, rules):
        report = evaluate(compliant, rules, moment=LAW_IN_FORCE)

        assert report.status is ContractStatus.GREEN, [
            (finding.code, finding.evidence) for finding in report.violations()
        ]

    def test_no_blocking_violations(self, compliant, rules):
        report = evaluate(compliant, rules, moment=LAW_IN_FORCE)

        assert report.blocking_violations() == []

    def test_key_rules_pass(self, compliant, rules):
        report = evaluate(compliant, rules, moment=LAW_IN_FORCE)

        for code in ("FTC-001", "AST-001", "ADR-001", "RTE-002", "TRV-001", "THR-001", "THR-002"):
            assert report.by_code(code).status is FindingStatus.PASSED, code

    def test_evidence_is_filled_for_passed_rules(self, compliant, rules):
        report = evaluate(compliant, rules, moment=LAW_IN_FORCE)

        passed = [f for f in report.findings if f.status is FindingStatus.PASSED]

        assert passed
        assert all(finding.evidence for finding in passed)


class TestViolatingContract:
    def test_status_is_red(self, violating, rules):
        report = evaluate(violating, rules, moment=LAW_IN_FORCE)

        assert report.status is ContractStatus.RED

    def test_qualification_missing(self, violating, rules):
        report = evaluate(violating, rules, moment=LAW_IN_FORCE)
        finding = report.by_code("FTC-001")

        assert finding.status is FindingStatus.FAILED
        assert "нерезидент" in finding.evidence

    def test_raw_wallet_is_reported_with_clause_number(self, violating, rules):
        report = evaluate(violating, rules, moment=LAW_IN_FORCE)
        finding = report.by_code("ADR-001")

        assert finding.status is FindingStatus.FAILED
        assert "TQn9Y2khEsLJW1ChVWFMSMeRDow5KcbLSE" in finding.evidence
        assert "2.3" in finding.clauses

    def test_asset_not_identified(self, violating, rules):
        report = evaluate(violating, rules, moment=LAW_IN_FORCE)
        finding = report.by_code("AST-001")

        assert finding.status is FindingStatus.FAILED
        assert "тикер" in finding.evidence

    def test_wrong_execution_moment(self, violating, rules):
        report = evaluate(violating, rules, moment=LAW_IN_FORCE)

        assert report.by_code("RTE-002").status is FindingStatus.FAILED

    def test_bank_registration_missing_above_threshold(self, violating, rules):
        report = evaluate(violating, rules, moment=LAW_IN_FORCE)
        finding = report.by_code("THR-002")

        assert finding.status is FindingStatus.FAILED
        assert "12 000 000" in finding.evidence

    def test_violation_count_is_substantial(self, violating, rules):
        report = evaluate(violating, rules, moment=LAW_IN_FORCE)

        assert len(report.blocking_violations()) >= 8


class TestApplicability:
    def test_deferred_rule_does_not_count_as_violation(self, violating, rules):
        report = evaluate(violating, rules, moment=LAW_IN_FORCE)
        finding = report.by_code("ADR-005")

        assert finding.status is FindingStatus.DEFERRED
        assert not finding.is_violation
        assert "2027-07-01" in finding.evidence

    def test_deferred_rule_activates_later(self, violating, rules):
        report = evaluate(violating, rules, moment=date(2027, 7, 1))

        assert report.by_code("ADR-005").status is not FindingStatus.DEFERRED

    def test_rule_without_date_stays_deferred(self, violating, rules):
        """Требование о репатриации не установлено, дата отсутствует."""
        report = evaluate(violating, rules, moment=date(2030, 1, 1))
        finding = report.by_code("RPT-003")

        assert finding.status is FindingStatus.DEFERRED
        assert "подзаконный акт" in finding.evidence

    def test_amount_is_formatted_without_breaking_sentence(self, violating, rules):
        """Разделитель разрядов не должен вырезать запятые из текста пояснения."""
        report = evaluate(violating, rules, moment=LAW_IN_FORCE)
        evidence = report.by_code("THR-001").evidence

        assert "12 000 000 руб." in evidence
        assert "№ 115-ФЗ, но обязанность" in evidence
        assert "  " not in evidence

    def test_manual_rules_are_empty_on_complete_sample(self, compliant, rules):
        report = evaluate(compliant, rules, moment=LAW_IN_FORCE)

        assert report.needs_manual_review() == []
        assert report.by_code("AML-002").status is FindingStatus.PASSED
        assert report.by_code("THR-003").status is FindingStatus.PASSED

    def test_splitting_schedule_is_flagged(self, rules, tmp_path: Path):
        from docx import Document as DocxDocument

        path = tmp_path / "split.docx"
        document = DocxDocument()
        document.add_paragraph(
            "1. Общая стоимость Товара составляет 12 000 000 рублей. "
            "Оплата: первый платёж 2 900 000 рублей, второй платёж 2 900 000 рублей, "
            "третий платёж 2 900 000 рублей, четвёртый платёж 3 300 000 рублей."
        )
        document.save(str(path))

        report = evaluate(ContractView.from_file(path), rules, moment=LAW_IN_FORCE)

        assert report.by_code("THR-003").status is FindingStatus.FAILED

    def test_agent_without_principal_fails(self, rules, tmp_path: Path):
        from docx import Document as DocxDocument

        path = tmp_path / "agent.docx"
        document = DocxDocument()
        document.add_paragraph("1. Покупатель действует как агент.")
        document.save(str(path))

        report = evaluate(ContractView.from_file(path), rules, moment=LAW_IN_FORCE)

        assert report.by_code("FTC-004").status is FindingStatus.FAILED

    def test_threshold_rule_skipped_below_amount(self, rules, tmp_path: Path):
        from docx import Document as DocxDocument

        small = tmp_path / "small.docx"
        document = DocxDocument()
        document.add_paragraph("1. Общая стоимость Товара составляет 50 000 рублей.")
        document.save(str(small))

        report = evaluate(ContractView.from_file(small), rules, moment=LAW_IN_FORCE)

        assert report.by_code("THR-002").status is FindingStatus.NOT_APPLICABLE
        assert report.by_code("TRV-002").status is FindingStatus.NOT_APPLICABLE


def test_report_separates_severity_levels(violating, rules):
    report = evaluate(violating, rules, moment=LAW_IN_FORCE)

    assert all(f.severity is Severity.MANDATORY for f in report.blocking_violations())
    assert all(f.severity is Severity.ADVISORY for f in report.advisory_violations())
    assert len(report.violations()) == len(report.blocking_violations()) + len(
        report.advisory_violations()
    )
